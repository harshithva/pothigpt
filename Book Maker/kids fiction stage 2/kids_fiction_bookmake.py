import openai
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re

# Set your OpenAI API key
openai.api_key = "os.getenv("OPENAI_API_KEY", "").strip()"

# Input and Output Paths
INPUT_PATH = "/Users/kuldeepsharma/Desktop/projectcode/Excel/kids_fiction_output.xlsx"
OUTPUT_DIR = "/Users/kuldeepsharma/Desktop/projectcode/Kids Fiction Books"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def add_hyperlink(paragraph, url, text):
    """
    Helper to add a hyperlink to a Word paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url, 
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    b.set(qn('w:val'), 'true')  
    rPr.append(b)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 24 Half-points = 12 pt font
    rPr.append(sz)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    text_el = OxmlElement('w:t')
    text_el.text = text

    new_run.append(rPr)
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def clean(text):
    """
    Basic cleanup to remove stray characters from text.
    """
    return re.sub(r'[\*#"]', '', str(text))

def generate_text(prompt):
    """
    Uses the OpenAI ChatCompletion endpoint to generate text from a prompt.
    Adapted for children's content with appropriate token limits.
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",  # same engine as earlier
            messages=[
                {"role": "system", "content": "You are a helpful assistant who writes engaging children's stories."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=4000,  # Reduced for children's content
            temperature=0.7
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        print(f"Error generating text: {e}")
        return "[Error generating content]"

def create_docx(book_title, chapters):
    """
    Creates a DOCX file for a given kids book_title and its associated chapters DataFrame.
    Adapted for children's books with age-appropriate formatting and content.
    """
    print(f"Creating kids DOCX for '{book_title}'")
    doc = Document()

    # Set the base style for Normal text - larger font for children
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)  # Larger font for children

    # Attempt to read the author's name from the first row
    author_name = ""
    if len(chapters) > 0:
        author_name = str(chapters.iloc[0].get("Author Name", "")).strip()

    # Book Title (level=0, centered) - larger font for children
    title_heading = doc.add_heading(book_title, 0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_heading.runs[0].font.size = Pt(24)  # Larger title font

    # "By <Author Name>" in 14pt, centered
    publishing_para = doc.add_paragraph()
    publishing_run = publishing_para.add_run(f"By {author_name}")
    publishing_run.font.size = Pt(14)
    publishing_run.font.name = "Calibri"
    publishing_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Page break before Copyright
    doc.add_page_break()

    # 1) Copyright Page - child-friendly version
    doc.add_heading('Copyright', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Copyright © 2025 AI Book Generator\n"
        "This is a work of fiction for children. Names, characters, places, and incidents either "
        "are the product of the author's imagination or are used fictitiously.\n"
        "Any resemblance to actual events, locales, or persons, living or dead, is "
        "entirely coincidental.\n"
        "All rights reserved.\n"
        "For permissions, contact support@yourplatform.com\n\n"
        "Recommended for ages 6-12"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 2) Prologue (if any text is in the 'Prologue' column of the first row)
    prologue_prompt = chapters.iloc[0]['Prologue'] if isinstance(chapters.iloc[0].get('Prologue', ''), str) else ""
    if prologue_prompt.strip():
        prologue_text = generate_text(prologue_prompt)
        # Remove heading instruction; start directly with the text
        doc.add_heading("Prologue", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(prologue_text)
        doc.add_page_break()

    # 3) Chapters & Epilogue
    for idx, row in chapters.iterrows():
        chap_title = str(row.get('Chapter', '')).strip()
        chap_prompt = row.get('Chapter Prompt', '')

        # If there's no chapter title or prompt, skip
        if not chap_title or not chap_prompt:
            continue

        # If it's an epilogue, handle final piece
        if chap_title.lower().startswith("epilogue"):
            epilogue_text = generate_text(chap_prompt)
            doc.add_heading("Epilogue", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(epilogue_text)
            doc.add_page_break()
        else:
            # Normal chapter: generate chapter text using the chapter prompt
            generated_chap_text = generate_text(chap_prompt)
            doc.add_heading(clean(chap_title), level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(generated_chap_text)
            doc.add_page_break()

    # Save the DOCX
    filename = os.path.join(OUTPUT_DIR, f"{book_title}.docx")
    doc.save(filename)
    print(f"Saved kids book: {filename}")

def process_books():
    # Read the Excel file
    df = pd.read_excel(INPUT_PATH)
    
    # Group by Book Title to get each set of rows for that book
    grouped = df.groupby("Book Title", sort=False)
    
    for book_title, chapters in grouped:
        # Ignore rows where "Chapter" is empty or NaN.
        chapters_filtered = chapters.dropna(subset=["Chapter"])
        chapters_filtered = chapters_filtered[chapters_filtered["Chapter"].str.strip() != ""]
        
        # Now create the DOCX only with these filtered chapter rows
        create_docx(book_title, chapters_filtered)

if __name__ == "__main__":
    all_titles = []

    # Read the Excel file
    df = pd.read_excel(INPUT_PATH)

    # Group by Book Title to get each set of rows for that book
    grouped = df.groupby("Book Title", sort=False)

    for book_title, chapters in grouped:
        # Ignore rows where "Chapter" is empty or NaN.
        chapters_filtered = chapters.dropna(subset=["Chapter"])
        chapters_filtered = chapters_filtered[chapters_filtered["Chapter"].str.strip() != ""]

        # Now create the DOCX only with these filtered chapter rows
        create_docx(book_title, chapters_filtered)

        # Get author name (from first row in group)
        author_name = str(chapters_filtered.iloc[0].get("Author Name", "")).strip()
        all_titles.append({"Book Title": book_title, "Author Name": author_name})

    # Save book-author list to Excel
    summary_df = pd.DataFrame(all_titles)
    summary_path = os.path.join(OUTPUT_DIR, "Kids_Book_Author_List.xlsx")
    summary_df.to_excel(summary_path, index=False)
    print(f"📘 Saved kids book-author list to: {summary_path}")

