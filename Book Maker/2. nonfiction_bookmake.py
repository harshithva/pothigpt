import os
import re
import openai
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── CONFIGURATION ──────────────────────────────────────────────────────────────

openai.api_key = os.getenv(
    "OPENAI_API_KEY",
    "os.getenv("OPENAI_API_KEY", "").strip()"
).strip()

MODEL = "gpt-4o-mini"
PROMPTS_EXCEL = "/Users/kuldeepsharma/Desktop/projectcode/Book_Generated_Content.xlsx"
WORD_OUTPUT_DIR = "/Users/kuldeepsharma/Desktop/projectcode/WordOutput"

# ── HELPERS ────────────────────────────────────────────────────────────────────

def format_text(text: str) -> str:
    return re.sub(r"[\*#\"]", "", text or "").strip()

def clean_intro(text: str, chapter_title: str) -> str:
    lines = [line.strip() for line in text.strip().split("\n") if line.strip()]
    cleaned = []
    subtitle = chapter_title.split(":")[-1].strip().lower()
    full_title = chapter_title.strip().lower()

    for line in lines:
        l = line.lower()
        if subtitle in l or full_title in l or re.match(r'chapter\s*\d+', l):
            continue
        cleaned.append(line)
    return "\n".join(cleaned).strip()

def clean_subsection(text: str, subheading: str) -> str:
    lines = [line.strip() for line in text.strip().split("\n") if line.strip()]
    cleaned = []
    sub_name = subheading.strip().lower()

    for line in lines:
        if sub_name in line.lower():
            continue
        cleaned.append(line)
    return "\n".join(cleaned).strip()

def generate_text(prompt: str) -> str:
    if not isinstance(prompt, str) or not prompt.strip():
        return ""
    try:
        print("🧠 Generating:", prompt[:60].strip().replace('\n', ' ') + "...")
        response = openai.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("❌ OpenAI error with prompt:", prompt[:60], "\n→", e)
        return ""

def init_doc(title: str) -> Document:
    doc = Document()
    os.makedirs(WORD_OUTPUT_DIR, exist_ok=True)

    # Title Page
    title_heading = doc.add_heading(title, level=0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_heading.runs[0].font.size = Pt(20)

    author_para = doc.add_paragraph("By AI Book Generator")
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.runs[0].font.size = Pt(14)
    doc.add_page_break()
    return doc

def save_doc(title: str, doc: Document):
    path = os.path.join(WORD_OUTPUT_DIR, f"{title}.docx")
    doc.save(path)
    print(f"💾 Saved progress to: {path}")

# ── MAIN PIPELINE ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    df = pd.read_excel(PROMPTS_EXCEL)
    books = {}

    for _, row in df.iterrows():
        title = row['Book_Title']
        if title not in books:
            print(f"\n📘 Starting new book: {title}")
            doc = init_doc(title)
            books[title] = doc

            if isinstance(row['Intro_Prompt'], str) and row['Intro_Prompt'].strip():
                print(f"📄 Generating Introduction for: {title}")
                intro_text = generate_text(row['Intro_Prompt'])
                if intro_text.strip():
                    doc.add_heading("Introduction", level=1)
                    doc.add_paragraph(format_text(intro_text))
                    doc.add_paragraph("")  # spacing
                    doc.add_page_break()
                    save_doc(title, doc)

        doc = books[title]
        chapter_title = row['Chapter_Title']
        print(f"📚 Chapter: {chapter_title}")
        chapter_intro = generate_text(row['Chapter_Intro'])
        chapter_intro = clean_intro(chapter_intro, chapter_title)

        doc.add_heading(chapter_title, level=1)
        doc.add_paragraph(format_text(chapter_intro))
        doc.add_paragraph("")

        for i in range(1, 5):
            sub_prompt = row.get(f'Subheading_{i}_Prompt')
            sub_title = row.get(f'Subheading_{i}')
            if isinstance(sub_prompt, str) and isinstance(sub_title, str):
                if sub_prompt.strip() and sub_title.strip():
                    print(f"   🔹 Subheading: {sub_title}")
                    sub_text = generate_text(sub_prompt)
                    sub_text = clean_subsection(sub_text, sub_title)
                    doc.add_heading(sub_title.strip(), level=2)
                    doc.add_paragraph(format_text(sub_text))
                    doc.add_paragraph("")

        doc.add_page_break()
        save_doc(title, doc)
